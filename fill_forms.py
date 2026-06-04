"""
填充实训手册中16张学习活动记录表
"""
from docx import Document
from docx.shared import Pt
import os

doc_path = r'C:\Users\DOVE\python学习\大二下实训\《深度学习应用开发综合实训》-学生手册-2026新版.docx'
doc = Document(doc_path)

# ======== 基本信息 ========
INFO = {
    '班级': '智能2412',
    '姓名': '胡世毅',
    '学号': '2405021228',
    '组别': '第X组',
    '组长': '（组长姓名）',
    '时间': '2026.5',
}

# ======== 16张表的内容 ========
RECORDS = [
    # 表1：总体方案设计与需求分析
    {
        '任务描述': (
            '1. 分析综合实训项目的四个应用场景需求：花卉图像分类、Titanic生存预测、'
            '时尚服饰图像分类、非线性系统回归预测。\n'
            '2. 确定系统功能模块：用户认证、模型推理、结果保存、历史查看、管理面板。\n'
            '3. 确定技术选型：前端Streamlit + 后端PyTorch + 数据库MySQL。\n'
            '4. 设计系统总体架构：backend（模型层）/ streamlist（展示层）/ dataset（数据层）三层架构。'
        ),
        '实际完成': (
            '1. 完成了四个应用场景的详细需求分析文档。\n'
            '2. 明确了每个场景的功能需求和数据流：用户输入→模型推理→结果展示→数据库保存。\n'
            '3. 确定了技术栈：Python 3.13 + PyTorch + Streamlit + MySQL + Pillow。\n'
            '4. 设计了项目目录结构：backend/flower_model, titanic_model, fashion_model, regression_model。'
        ),
        '问题': (
            '1. 四个场景的技术难度差异大，花卉分类是图像分类、Titanic是表格数据、回归是时序预测。\n'
            '2. 如何统一四个不同输入模态（图像/表格/序列）的Web界面设计。\n\n'
            '解决方法：采用Streamlit的灵活组件体系，每个场景独立设计表单和结果展示区域，'
            '共享侧边栏导航和数据库接口。'
        ),
        '小结': (
            '通过需求分析和总体设计，明确了项目的技术路线和功能边界。三层架构设计'
            '（模型层/展示层/数据层）实现了关注点分离，为后续并行开发打下了基础。'
            '学习了软件工程中需求分析和系统设计的基本方法。'
        ),
    },
    # 表2：系统功能设计与技术选型
    {
        '任务描述': (
            '1. 细化系统功能模块设计：主页展示、四个AI预测模块、用户系统、管理面板。\n'
            '2. 确定每个模块的UI布局方案。\n'
            '3. 设计用户认证流程：注册→登录→session管理→权限控制。\n'
            '4. 设计数据库ER图和表结构。'
        ),
        '实际完成': (
            '1. 主页设计：四张场景卡片 + 系统特色 + 快速开始引导。\n'
            '2. 用户系统设计：users/user_activities/model_usage_stats/image_classifications 四张数据表。\n'
            '3. 设计了@require_auth装饰器实现统一鉴权。\n'
            '4. 每个模型模块设计为"预测标签页+历史记录标签页+统计分析标签页"三标签结构。\n'
            '5. 管理员面板包含：用户管理、活动记录、模型统计、系统统计、图像记录管理五个功能。'
        ),
        '问题': (
            '1. Streamlit是无状态框架，如何在页面刷新间保持登录状态。\n'
            '2. 管理员和普通用户的权限如何优雅分离。\n\n'
            '解决方法：使用st.session_state保存登录信息；通过user_role字段配合条件渲染'
            '实现基于角色的访问控制（RBAC）。'
        ),
        '小结': (
            '掌握了Web应用系统功能模块的拆解方法。理解了无状态框架中的状态管理机制'
            '（st.session_state）。学会了数据库ER图绘制和表结构设计。RBAC设计模式简单实用。'
        ),
    },
    # 表3：花卉分类数据集准备
    {
        '任务描述': (
            '1. 收集花卉图像数据集，包含5个类别：雏菊(daisy)、蒲公英(dandelion)、'
            '玫瑰(rose)、向日葵(sunflower)、郁金香(tulip)。\n'
            '2. 分析数据集规模、类别分布、图像质量。\n'
            '3. 设计数据预处理流水线：Resize→数据增强→ToTensor→Normalize。\n'
            '4. 划分训练集(70%)、验证集(15%)、测试集(15%)。'
        ),
        '实际完成': (
            '1. 花卉数据集存放于dataset/flowers/，5个类别共约5000+张图片。\n'
            '2. 训练集数据增强：RandomHorizontalFlip(0.5), RandomRotation(20°), '
            'RandomAffine(shear=10, scale=0.8-1.2), ColorJitter(0.2)。\n'
            '3. 验证/测试集仅做Resize(224,224) + ToTensor + Normalize。\n'
            '4. 采用sklearn的train_test_split进行分层抽样划分，保证各类别比例一致。\n'
            '5. 创建了TransformedDataset类使每个子集独立持有transform，避免共享覆盖问题。'
        ),
        '问题': (
            '1. 初始代码使用random_split后直接修改val_dataset.dataset.transform，'
            '导致训练集的transform被覆盖，数据增强完全失效。\n\n'
            '解决方法：自定义TransformedDataset类，加载ImageFolder时不带transform，'
            '每个子集在__getitem__中独立调用自己的transform，彻底避免了共享状态问题。'
        ),
        '小结': (
            '深入理解了PyTorch Dataset/DataLoader体系的工作原理。random_split返回的'
            'Subset对象共享底层Dataset，直接修改transform会产生全局副作用。学会了'
            '通过自定义Dataset类来灵活控制数据预处理流程。数据增强是提高模型泛化能力的关键技术。'
        ),
    },
    # 表4：Titanic/Fashion/回归数据集准备
    {
        '任务描述': (
            '1. Titanic数据集：从seaborn加载，包含乘客信息（舱位、性别、年龄、票价等）。\n'
            '2. Fashion-MNIST：从本地IDX文件读取，10类服饰，28×28灰度图。\n'
            '3. 回归数据集：生成Lorenz-like混沌非线性系统合成数据。\n'
            '4. 各数据集进行清洗、编码、标准化、划分。'
        ),
        '实际完成': (
            '1. Titanic：特征工程（family_size, is_alone两个衍生特征），'
            '缺失值填充（中位数），Label Encoding（性别/港口），StandardScaler标准化。\n'
            '2. Fashion-MNIST：编写read_idx()函数解析IDX二进制格式（大端序），'
            '将numpy数组转为TensorDataset，按0.1比例分出验证集。\n'
            '3. 回归数据：用numpy生成了5000条非线性数据，公式包含sin/cos/exp/高阶项，验证MLP和LSTM对不同复杂度映射的学习能力。'
        ),
        '问题': (
            '1. Fashion-MNIST的IDX文件最初为空（0字节），torchvision在线下载失败（网络限制）。\n'
            '2. IDX文件嵌套在同名文件夹内，路径拼接需要额外处理。\n'
            '3. Windows下DataLoader的num_workers>0会导致多进程序列化报错。\n\n'
            '解决方法：手动获取数据后修正路径为dataset/cloths/文件名/文件名；'
            '将num_workers统一设为0，避免Windows multiprocessing的pickling问题。'
        ),
        '小结': (
            '掌握了多种数据加载方式：pandas读表格、seaborn数据集、IDX二进制格式解析、'
            '合成数据生成。理解了不同模型需要不同的数据格式：图像→4D张量(N,C,H,W)、'
            '表格→2D张量(N,features)、时序→3D张量(N,seq_len,features)。'
        ),
    },
    # 表5：花卉分类模型设计（ResNet50迁移学习）
    {
        '任务描述': (
            '1. 设计花卉分类的两种深度神经网络架构。\n'
            '架构A：ResNet50迁移学习——利用ImageNet预训练权重，替换FC层为Dropout(0.5)→'
            'Linear(2048→512)→ReLU→Dropout(0.3)→Linear(512→5)。\n'
            '2. 确定超参数：Adam(lr=0.001), CrossEntropyLoss, batch_size=32, epochs=50。\n'
            '3. 配置学习率调度ReduceLROnPlateau(patience=5, factor=0.5)和早停(patience=10)。'
        ),
        '实际完成': (
            '1. ResNet50模型定义：torchvision.models.resnet50(weights=IMAGENET1K_V1)加载预训练权重，'
            '替换fc层为自定义分类头（双层FC+Dropout）。\n'
            '2. 完善了训练函数：train_epoch（进度条tqdm+实时损失/准确率）+ validate（无梯度计算）。\n'
            '3. 实现了checkpoint保存：存储model_state_dict, optimizer_state_dict, val_acc, class_names。\n'
            '4. 训练曲线可视化：matplotlib绘制train/val的Loss和Accuracy双图。'
        ),
        '问题': (
            '1. ResNet50参数量大(~23M)，CPU训练速度慢。\n'
            '2. 数据加载使用num_workers=4在Windows上报错。\n\n'
            '解决方法：检测CUDA可用性自动切换device；Windows下改用num_workers=0。'
        ),
        '小结': (
            '深入理解了迁移学习的原理和优势：利用ImageNet预训练的底层特征提取能力，'
            '只需微调顶部分类层即可适配新任务，大幅降低了对标注数据量和训练时间的需求。'
            '学会了使用ReduceLROnPlateau动态调整学习率和Early Stopping防止过拟合。'
        ),
    },
    # 表6：花卉分类SimpleCNN + 双模型合并保存
    {
        '任务描述': (
            '1. 设计SimpleCNN架构用于对比实验：4层卷积(3→64→128→256→512) + BN + ReLU + '
            'MaxPool + 自适应池化 + 2层FC。\n'
            '2. 修改训练脚本支持双模型训练并合并保存到单个checkpoint文件。\n'
            '3. 在Streamlit中实现模型架构切换功能。'
        ),
        '实际完成': (
            '1. SimpleCNN模型约5M参数，相比ResNet50(~23M)更轻量，适合从零训练。\n'
            '2. 训练脚本改造：依次训练ResNet50和SimpleCNN，分别保存最优state_dict。\n'
            '3. 合并checkpoint：单文件存储resnet50_state_dict和simple_cnn_state_dict及其精度。\n'
            '4. Streamlit界面增加模型选择下拉框：resnet50(迁移学习) vs simple_cnn(轻量级)。'
        ),
        '问题': (
            '1. 旧checkpoint格式只有model_state_dict，新代码期望resnet50_state_dict键名，会导致加载失败。\n'
            '2. SimpleCNN训练epochs需要比ResNet50更多（80 vs 50）才能收敛。\n\n'
            '解决方法：增加fallback逻辑；SimpleCNN采用更大的epochs和更长的早停patience。'
        ),
        '小结': (
            '通过对比ResNet50和SimpleCNN两种架构，深入理解了"预训练知识迁移"和'
            '"任务专用轻量网络"两种设计哲学的区别和适用场景。学会了将多个模型统一管理'
            '在单一checkpoint文件中的工程技巧。'
        ),
    },
    # 表7：Titanic生存预测模型设计
    {
        '任务描述': (
            '1. 设计Titanic生存预测的两种MLP架构。\n'
            '架构A（深度型）：Linear(9→128)→BN→ReLU→Dropout(0.4)→Linear(128→64)→BN→ReLU→'
            'Dropout(0.3)→Linear(64→32)→BN→ReLU→Dropout(0.2)→Linear(32→1)→Sigmoid。\n'
            '架构B（宽型）：Linear(9→256)→ReLU→Dropout(0.5)→Linear(256→128)→ReLU→'
            'Dropout(0.3)→Linear(128→1)→Sigmoid。\n'
            '2. 训练配置：BCELoss, Adam(lr=0.001, weight_decay=1e-4), early_stop=30。'
        ),
        '实际完成': (
            '1. 实现了TitanicMLP和TitanicMLP_Wide两个模型类。\n'
            '2. 特征工程：从7个原始特征扩展到9个（+family_size, +is_alone）。\n'
            '3. 训练流程：分层抽样(70/15/15) → StandardScaler → TensorDataset → DataLoader。\n'
            '4. checkpoint同时保存mlp和mlp_wide两个state_dict及scaler参数（mean, scale）。\n'
            '5. Streamlit界面：7项表单输入→实时推理→生存概率+特征影响柱状图。'
        ),
        '问题': (
            '1. Titanic数据集较小(~891条)，模型容易过拟合。\n'
            '2. MLP深度型和宽型在测试集上精度差异不大，如何选择最佳保存。\n\n'
            '解决方法：增加weight_decay正则化和Dropout防过拟合；分别测试两种架构后取'
            '验证集精度更高的作为默认加载模型，但两个权重都保存以便对比。'
        ),
        '小结': (
            '掌握了表格数据的深度学习处理流程：特征工程→标准化→MLP分类。理解了'
            'Batch Normalization加速收敛和Dropout防止过拟合的原理。对比了"深度+BN"'
            '和"宽度+Dropout"两种MLP设计策略的优缺点。'
        ),
    },
    # 表8：时尚服饰分类模型设计
    {
        '任务描述': (
            '1. 设计10类时尚服饰识别的两种CNN架构。\n'
            '架构A（FashionCNN）：7层卷积(1→32→32→64→64→128→128) + BN + Dropout + 全局池化 + 2层FC。\n'
            '架构B（FashionResNet）：ResNet18迁移学习，修改第一层conv1为单通道输入，替换fc为10分类。\n'
            '2. 适配28×28灰度图，标准化参数为单通道(mean=0.2860, std=0.3530)。'
        ),
        '实际完成': (
            '1. FashionCNN：手工设计的轻量CNN，Conv2d(1,32,3)→...→AdaptiveAvgPool→Linear(128,256)→'
            'Dropout(0.4)→Linear(256,10)。\n'
            '2. FashionResNet：替换conv1 = Conv2d(1,64,kernel=7)，fc = Linear(512,10)。\n'
            '3. 训练配置：CrossEntropyLoss, Adam, CosineAnnealingLR, batch_size=128, epochs=30。\n'
            '4. 本地IDX数据加载：read_idx()解析二进制→numpy→TensorDataset→TransformDataset包装标准化。'
        ),
        '问题': (
            '1. torchvision.datasets.FashionMNIST下载失败（网络限制），需要手动加载本地IDX文件。\n'
            '2. IDX文件最初显示0字节（Git LFS占位符），实际数据嵌套在同名子文件夹中。\n'
            '3. 自定义TransformDataset作为DataLoader的局部类导致pickle错误。\n\n'
            '解决方法：编写read_idx函数直接解析IDX大端序二进制；修正路径嵌套问题；'
            '将TransformDataset提升为模块级类，num_workers=0避免Windows多进程问题。'
        ),
        '小结': (
            '学会了处理IDX格式的二进制数据文件。理解了CNN处理灰度图与RGB图的区别'
            '（输入通道数和标准化参数）。掌握了余弦退火学习率调度策略（CosineAnnealingLR）。'
            '自定义Dataset+Transform包装是PyTorch数据加载的灵活模式。'
        ),
    },
    # 表9：非线性系统回归模型设计
    {
        '任务描述': (
            '1. 设计非线性回归的两种架构。\n'
            '架构A（RegressionMLP）：4层全连接(3→128→128→64→3)，Dropout(0.2)，无激活输出层。\n'
            '架构B（RegressionLSTM）：2层LSTM(hidden=64, dropout=0.2) + Linear(64→1)，处理时序。\n'
            '2. 生成合成非线性数据：Lorenz-like混沌函数，包含sin/cos/exp/高阶耦合项+噪声。\n'
            '3. 评估指标：R²（决定系数），目标≥0.50。'
        ),
        '实际完成': (
            '1. MLP回归器：3维输入→3维输出，MSELoss，训练200 epochs。\n'
            '2. LSTM回归器：10步时序窗口→1维输出，输入形状(batch, seq_len=10, features=3)。\n'
            '3. 数据合成：x₁,x₂,x₃∈[-3,3]，y = sin(1.5x₁)·cos(0.8x₂) + 0.3x₃² + '
            '0.5exp(-|x₂|) + 0.4sin(0.3x₁x₂) + noise。\n'
            '4. Streamlit界面：MLP模式输入3维特征→预测3维输出；LSTM模式输入10行序列→预测下一时刻。'
        ),
        '问题': (
            '1. MLP和LSTM数据格式和标准化方式不同（MLP用StandardScaler on features，'
            'LSTM需per-feature flatten后再标准化）。\n'
            '2. LSTM的序列数据生成需要使用滑动窗口，要保证不泄露未来信息。\n\n'
            '解决方法：分开处理MLP和LSTM数据，分别保存各自的scaler；'
            '滑动窗口用for循环从前到后依次构造，确保每个窗口的target是紧接窗口之后的值。'
        ),
        '小结': (
            '理解了非线性系统的"黑箱"建模方法。MLP适合静态特征→输出映射，'
            'LSTM适合带有时序依赖的动态系统预测。R²是回归模型的重要评估指标，'
            '衡量预测值与真实值的拟合优度。学会了滑动窗口方法构造时序监督学习数据。'
        ),
    },
    # 表10：数据库设计与实现
    {
        '任务描述': (
            '1. 设计MySQL数据库系统，支持多用户、活动日志、模型统计和图像存储。\n'
            '2. 设计E-R图：用户(User)→活动(Activity)、用户→图像分类记录(ImageClassification)。\n'
            '3. 编写数据库初始化脚本，创建表和默认数据。\n'
            '4. 实现数据库CRUD操作的Python接口。'
        ),
        '实际完成': (
            '1. 创建了4张数据表：\n'
            '  - users：id, username, password_hash, email, role, is_active, last_login\n'
            '  - user_activities：外键user_id, activity_type, model_name, confidence_score\n'
            '  - model_usage_stats：model_name(UNIQUE), total_predictions, avg_confidence\n'
            '  - image_classifications：user_id(FK), image_data(LONGBLOB), class_probabilities(TEXT/JSON)\n'
            '2. 索引设计：username, user_id, created_at, predicted_class, confidence_score。\n'
            '3. 外键约束：ON DELETE CASCADE级联删除。\n'
            '4. 默认用户：admin/admin123, demo_user/demo123。'
        ),
        '问题': (
            '1. 图像以LONGBLOB存储在数据库中导致表体积快速膨胀。\n'
            '2. 初始设计中auth.py在模块导入时就执行init_database()，产生副作用。\n\n'
            '解决方法：上传前thumbnail压缩至800×800并quality=85；'
            '将init_database()调用从模块级别移到app.py的main()函数中显式调用。'
        ),
        '小结': (
            '掌握了MySQL数据库的完整设计流程：需求分析→ER图→建表SQL→索引优化→Python接口。'
            '理解了外键约束和级联操作的作用。学会了使用pymysql和SQLAlchemy两种方式操作数据库。'
            '数据库初始化的模块级调用是一个常见的反模式，应显式调用而非依赖import副作用。'
        ),
    },
    # 表11：用户认证系统实现
    {
        '任务描述': (
            '1. 实现用户注册/登录/密码修改功能。\n'
            '2. 实现SHA-256密码哈希存储。\n'
            '3. 实现session管理和登录状态保持。\n'
            '4. 实现角色访问控制（RBAC）：user和admin权限分离。\n'
            '5. 实现活动日志记录系统。'
        ),
        '实际完成': (
            '1. 注册功能：用户名唯一性检查 + 邮箱格式验证 + 密码强度验证(≥6位,含字母+数字)。\n'
            '2. 登录功能：SHA-256哈希比对 → 更新last_login → 记录login活动 → 写入session_state。\n'
            '3. @require_auth装饰器：检查logged_in状态，未登录时自动显示登录界面。\n'
            '4. admin面板：用户管理(查看/删除)、活动记录(筛选/CSV导出)、模型统计(柱状图)、'
            '图像记录管理(筛选/分页/预览/删除)。\n'
            '5. 活动日志：login/logout/prediction/register/change_password五种类型。'
        ),
        '问题': (
            '1. Streamlit每次交互重新运行整个脚本，如何保持登录状态。\n'
            '2. 数据库连接的资源管理（每次操作都要get_connection + close）。\n\n'
            '解决方法：使用st.session_state在重新运行间保持用户信息；'
            '每个数据库操作采用try-finally确保连接正确关闭。'
        ),
        '小结': (
            '掌握了Web应用用户认证系统的核心实现：密码哈希、session管理、'
            '装饰器鉴权、RBAC权限控制。理解了Streamlit无状态框架中session_state的作用。'
            '活动日志系统为后续用户行为分析和系统优化提供了数据基础。'
        ),
    },
    # 表12：Streamlit界面设计与实现
    {
        '任务描述': (
            '1. 设计并实现Streamlit Web应用的整体界面布局。\n'
            '2. 主页：系统标题、四张场景卡片、系统特色、快速开始引导。\n'
            '3. 侧边栏：导航菜单、用户信息、系统信息。\n'
            '4. 自定义CSS样式美化界面。\n'
            '5. 四个模型模块的完整UI实现。'
        ),
        '实际完成': (
            '1. 页面配置：title="AI综合应用系统", layout="wide", icon="🤖"。\n'
            '2. 主页：渐变色标题(background-clip: text)、4张悬停浮起卡片(hover translateY)。\n'
            '3. 侧边栏：radio导航、用户头像/角色/邮箱/last_login信息展示。\n'
            '4. 每个模型模块采用tabs标签页(预测/历史/统计)，col1/col2双栏布局。\n'
            '5. 模型架构选择下拉框 + 实时推理 + 概率柱状图 + 置信度仪表盘。\n'
            '6. 更新所有use_container_width→width="stretch"（Streamlit弃用API迁移）。'
        ),
        '问题': (
            '1. Streamlit表单提交后结果消失（因为页面重新运行但未保存状态）。\n'
            '2. 按钮点击可能导致重复提交。\n'
            '3. st.dataframe列名中文显示问题。\n\n'
            '解决方法：用st.form包裹输入+st.form_submit_button合并提交；'
            '每次成功预测后将结果存入session_state；适当使用st.rerun()刷新界面。'
        ),
        '小结': (
            '学会了使用Streamlit快速构建数据科学Web应用。掌握了st.columns布局、'
            'st.tabs标签页、st.expander折叠面板、st.metric指标卡片等组件的使用方法。'
            '自定义CSS增强了界面美观度。模型懒加载缓存机制避免重复加载大文件。'
        ),
    },
    # 表13：开发环境搭建与配置
    {
        '任务描述': (
            '1. 配置Python 3.13开发环境，安装PyTorch、Streamlit等依赖。\n'
            '2. 安装和配置MySQL数据库服务。\n'
            '3. 配置项目目录结构和Python路径。\n'
            '4. 编写setup.bat和start.bat实现一键安装和启动。'
        ),
        '实际完成': (
            '1. 环境配置：Python 3.13 + PyTorch + torchvision + Streamlit + '
            'pymysql + pandas + numpy + Pillow + scikit-learn + matplotlib + tqdm + seaborn。\n'
            '2. setup.bat：自动检查Python环境 → 安装依赖（清华镜像加速）→ 初始化数据库。\n'
            '3. start.bat：检查Python → 初始化DB → 启动Streamlit → 自动打开浏览器。\n'
            '4. 路径管理：使用os.path.abspath(__file__)获取脚本绝对路径，动态计算project_root。'
        ),
        '问题': (
            '1. Windows CMD对UTF-8编码的bat文件支持差，出现乱码和命令截断。\n'
            '2. start.bat中--server.headless true导致不自动打开浏览器。\n'
            '3. 不同机器Python路径不同，硬编码导致bat无法运行。\n\n'
            '解决方法：去掉chcp 65001和中文注释，改用纯英文bat；移除headless参数；'
            '在bat开头用set PYTHON=...变量，方便用户修改。'
        ),
        '小结': (
            '掌握了Python项目环境管理的工程实践。理解了Windows bat脚本的编码限制'
            '（应使用系统ANSI编码而非UTF-8）。学会了编写自动化脚本简化开发和部署流程。'
            '项目路径管理应使用绝对路径而非依赖当前工作目录。'
        ),
    },
    # 表14：系统集成与功能联调
    {
        '任务描述': (
            '1. 将四个深度学习模型集成到Streamlit前端。\n'
            '2. 实现模型推理与数据库保存的联动。\n'
            '3. 实现历史记录查看和统计分析功能。\n'
            '4. 端到端测试每个场景的完整流程。\n'
            '5. 修复集成过程中发现的bug。'
        ),
        '实际完成': (
            '1. 模型加载：实现_model_cache懒加载单例，避免每次交互重载大模型。\n'
            '2. 推断→保存流程：用户输入→模型推理→log_activity记录→save_classified_image保存图像→'
            'update_model_stats更新统计。\n'
            '3. 历史记录：get_user_classifications查询→expander展示→可展开概率分布→删除按钮。\n'
            '4. 统计分析：各类别数量柱状图+平均置信度折线图+总/今日统计卡片。\n'
            '5. 修复的bug：重复导入清理、relative_path→absolute_path、'
            'datetime.now()→datetime.datetime.now()、use_container_width→width。'
        ),
        '问题': (
            '1. get_user_activities返回的记录中model_name字段为None，导致.startswith()报错。\n'
            '2. st.dataframe方法use_container_width参数被弃用，产生大量warning。\n'
            '3. 花卉模型checkpoint格式不兼容新旧两种结构。\n\n'
            '解决方法：用(a.get("model_name") or "")替代a.get("model_name", "")处理None值；'
            '全局替换width="stretch"；checkpoint加载增加兼容逻辑。'
        ),
        '小结': (
            '系统集成是发现和修复边界问题的重要阶段。懒加载缓存是Streamlit这类"每次交互重跑"'
            '框架的关键性能优化手段。数据库、前端和后端模型的交互中，错误处理和数据格式'
            '一致性是最容易出现问题的环节。'
        ),
    },
    # 表15：项目文档编写
    {
        '任务描述': (
            '1. 编写项目报告书：需求分析、系统设计、模型设计、数据库设计、界面设计、测试报告。\n'
            '2. 编写产品使用说明书：安装步骤、各模块使用指南、常见问题。\n'
            '3. 填写学生手册中的学习活动记录表。\n'
            '4. 准备项目答辩PPT。'
        ),
        '实际完成': (
            '1. 项目报告结构：9个章节（概述→需求分析→总体设计→模型设计→数据库设计→'
            '界面设计→实现→测试→总结）。\n'
            '2. 产品使用说明书：setup.bat→train models→start.bat三步启动流程。\n'
            '3. 学生手册：16张学习活动记录表，覆盖从需求分析到答辩的全过程。\n'
            '4. 整体代码约2000+行Python，9个.py源文件，4个checkpoint文件。'
        ),
        '问题': (
            '1. 文档编写工作量与代码开发并行的难度。\n'
            '2. 如何将技术细节用简洁清晰的语言表述。\n\n'
            '解决方法：采用"任务描述→实际完成→问题→小结"四段式结构，'
            '既记录工作过程也沉淀经验教训。在代码开发的同时维护文档。'
        ),
        '小结': (
            '文档编写是软件工程不可或缺的环节。好的文档让项目可维护、可传承。'
            '学会了项目报告、使用说明书、答辩PPT的规范写法。'
            '四段式记录法（任务/完成/问题/小结）是一种高效的实训日志记录方式。'
        ),
    },
    # 表16：项目总结与收获
    {
        '任务描述': (
            '1. 全面回顾整个综合实训项目的开发过程。\n'
            '2. 总结学到的知识点和技术能力。\n'
            '3. 分析项目中存在的不足和改进方向。\n'
            '4. 准备最终答辩演示。'
        ),
        '实际完成': (
            '1. 完成了一个完整的深度学习应用系统：4个AI场景×2种架构、用户认证、数据库管理、Web界面。\n'
            '2. 技术栈覆盖：PyTorch(模型定义/训练/推理) + Streamlit(Web UI) + MySQL(数据持久化) + '
            'NumPy/Pandas(数据处理)。\n'
            '3. 核心工程实践：三层架构设计、单文件双模型checkpoint、懒加载缓存、自动化bat脚本、'
            '弃用API迁移、Windows兼容性处理。\n'
            '4. 人均代码量约400行，总代码量约2000+行Python。'
        ),
        '问题': (
            '1. 花卉模型目前仅加载ResNet50权重，SimpleCNN的权重需要重新训练后合并保存。\n'
            '2. 三个"完成"模型的实际训练精度待验证确认。\n'
            '3. 系统缺少单元测试和自动化测试。\n\n'
            '改进方向：完善花卉双模型训练；添加pytest测试用例；使用Docker容器化部署；'
            '密码哈希升级为bcrypt。'
        ),
        '小结': (
            '本次综合实训全面锻炼了深度学习应用开发的全流程能力：从数据准备→模型设计→'
            '训练优化→Web集成→数据库管理→文档编写。最深的体会是"工程化思维"的重要性'
            '——好的代码不仅要能运行，还要结构清晰、易于维护、方便部署。团队协作中，'
            '明确的分工和良好的沟通是项目成功的关键。'
        ),
    },
]

# ======= 填充表格 =======
# 表5到表20对应记录1到16
for i, record in enumerate(RECORDS):
    table_idx = 5 + i  # Table 5 ~ Table 20
    if table_idx >= len(doc.tables):
        print(f'Table {table_idx} not found, skipping')
        break

    table = doc.tables[table_idx]

    # Row 0: 班级, (empty), 姓名, (empty), 学号, (empty)
    # Row 1: 组别, (empty), 组长, (empty), 时间, (empty)
    # Rows 2-5: 跨6列合并的大单元格
    # Row 6: 教师评阅

    try:
        # Row 0: 班级/姓名/学号
        table.cell(0, 1).text = INFO['班级']
        table.cell(0, 3).text = INFO['姓名']
        table.cell(0, 5).text = INFO['学号']

        # Row 1: 组别/组长/时间
        table.cell(1, 1).text = INFO['组别']
        table.cell(1, 3).text = INFO['组长']
        table.cell(1, 5).text = INFO['时间']

        # Row 2: 任务描述（填入标签右侧的空白格 cell(2,1)）
        table.cell(2, 1).text = record['任务描述']

        # Row 3: 实际完成内容
        table.cell(3, 1).text = record['实际完成']

        # Row 4: 遇到的问题及解决方法
        table.cell(4, 1).text = record['问题']

        # Row 5: 实训小结
        table.cell(5, 1).text = record['小结']

        # Row 6: 教师评阅（留空）

        print(f'Table {table_idx} (Record {i+1}): OK')
    except Exception as e:
        print(f'Table {table_idx} Error: {e}')

# 保存
import time
output_path = doc_path.replace('.docx', f'_filled_{int(time.time())}.docx')
doc.save(output_path)
print(f'\nSaved to: {output_path}')
